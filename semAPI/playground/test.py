from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mammoth

import imgkit

import sys
import os
import pypandoc
import docx




def add_footer():
    file = 'w.docx'
    
    document = Document(f'{file}')

    footer = document.sections[0].footer
    # Check if there is a footer already and remove it
    if footer.is_linked_to_previous is False:
        footer.is_linked_to_previous = True


    ftable= footer.add_table(1, 2, Inches(6))
    ftab_cells= ftable.rows[0].cells
    ft0= ftab_cells[1].add_paragraph()
    ft0.alignment = WD_ALIGN_PARAGRAPH.RIGHT 

    kh=ft0.add_run()
    kh.add_picture('1.png', width=Inches(1))
    
    ft0= ftab_cells[0].add_paragraph()
    ft0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kh=ft0.add_run()
    kh.add_picture('1.png', width=Inches(1))

    document.save('w.docx')

def convert_docx_to_pdf():
    file = 'w.docx'
    os.system(f'unoconv {file}')
    


add_footer()
convert_docx_to_pdf()
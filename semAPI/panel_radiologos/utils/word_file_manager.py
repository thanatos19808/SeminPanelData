from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
from docx import Document
from ..models import Imagenologia
from django.conf import settings
from PyPDF2 import PdfFileWriter, PdfFileReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io


class WordFileManager():
    def __init__(
        self,
        file_name: str,
    ):
        self.file_name = file_name
        
    
    def add_footer(
        self,
        imagenologo_id: int = None,
        ) -> bool:
        folder_for_water_marks = settings.WATER_MARKS    
        document = Document(f'{self.file_name}')

        footer = document.sections[0].footer
        # Check if there is a footer already and remove it
        if footer.is_linked_to_previous is False:
            footer.is_linked_to_previous = True


        ftable= footer.add_table(1, 2, Inches(6))
        ftab_cells= ftable.rows[0].cells
        ft0= ftab_cells[1].add_paragraph()
        ft0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        kh=ft0.add_run()
        kh.add_picture('media/footer.png', width=Inches(1))
        
        # Check if user has water mark
        user = Imagenologia.objects.get(
            id=imagenologo_id
        )
        if user.water_mark is not None:            
            ft0= ftab_cells[0].add_paragraph()
            ft0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            kh=ft0.add_run()
            kh.add_picture(f'{folder_for_water_marks}/{user.water_mark}', width=Inches(1))
        

        document.save(f'{self.file_name}')
        
        return True
    
    
    def convert_docx_to_pdf(self) -> bool:
        os.system(f'unoconv {self.file_name}')
        return True
    
    def remove_docx(self) -> bool:
        os.remove(self.file_name)
    
    def remove_pdf(
        self,
        user_folder: str = None,
        old_file: str = None
        ) -> bool:
        if user_folder:            
            dir_files = os.listdir(user_folder)
            for item in dir_files:
                if item.endswith(".pdf"):
                    os.remove(os.path.join(user_folder, item))
        else:
            os.remove(self.file_name)
        
    def add_water_mark_to_pdf(
        self,
        folder_container: str,
        imagenologo_id: int = None,
        ) -> bool:
        folder_for_water_marks = settings.WATER_MARKS
        packet = io.BytesIO()
        # create a new PDF with Reportlab
        user = Imagenologia.objects.get(
            id=imagenologo_id
        )
        can = canvas.Canvas(packet, pagesize=letter)
        can.drawImage('media/footer.png', 500, 2, 100, 45, mask='auto')
        if user.water_mark is not None: 
            can.drawImage(f'{folder_for_water_marks}/{user.water_mark}', 0, 2, 100, 45, mask='auto')
        can.save()
        
        # move to the beginning of the StringIO buffer
        packet.seek(0)
        new_pdf = PdfFileReader(packet)
        # read your existing PDF
        existing_pdf = PdfFileReader(open(self.file_name, 'rb'))
        output = PdfFileWriter()
        # add the "watermark" (which is the new pdf) on the existing page
        for i in range(existing_pdf.getNumPages()):
            page = existing_pdf.getPage(i)
            page.mergePage(new_pdf.getPage(0))
            output.addPage(page)
        # finally, write "output" to a real file
        outputStream = open(f'{folder_container}/interpretacion.pdf', 'wb')
        output.write(outputStream)
        outputStream.close()
        
    